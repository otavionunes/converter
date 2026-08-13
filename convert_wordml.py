#!/usr/bin/env python3
"""
convert_wordml.py — Convert Word 2003 XML (.doc) to .docx using lxml + python-docx.
No LibreOffice needed. ~1s for 2MB files.
Usage: python3 convert_wordml.py <input.doc> <output.docx>
"""
import sys
import os


def convert(input_path, output_path):
    try:
        import lxml.etree as ET
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        print(f"ERROR: missing dependency: {e}", file=sys.stderr)
        sys.exit(4)

    W = 'http://schemas.microsoft.com/office/word/2003/wordml'

    def w(tag): return f'{{{W}}}{tag}'

    with open(input_path, 'rb') as f:
        raw = f.read()

    if b'encoding=' not in raw[:100]:
        raw = raw.replace(b'<?xml version="1.0"?>', b'<?xml version="1.0" encoding="UTF-8"?>', 1)

    tree = ET.fromstring(raw)
    doc = Document()

    # Remove default empty paragraph
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    def get_run_text(elem):
        return ''.join(t.text or '' for t in elem.findall(f'.//{w("t")}'))

    def add_para(wp):
        p = doc.add_paragraph()
        ppr = wp.find(w('pPr'))
        if ppr is not None:
            jc = ppr.find(w('jc'))
            if jc is not None:
                val = jc.get(w('val'), '')
                if val == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif val == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pstyle = ppr.find(w('pStyle'))
            if pstyle is not None:
                sval = pstyle.get(w('val'), '')
                if 'heading' in sval.lower() or 'Heading' in sval:
                    try:
                        level = int(''.join(c for c in sval if c.isdigit()) or '1')
                        p.style = f'Heading {min(level, 9)}'
                    except (ValueError, KeyError):
                        pass
        for wr in wp.findall(w('r')):
            text = get_run_text(wr)
            if not text:
                continue
            run = p.add_run(text)
            rpr = wr.find(w('rPr'))
            if rpr is not None:
                if rpr.find(w('b')) is not None: run.bold = True
                if rpr.find(w('i')) is not None: run.italic = True
                sz = rpr.find(w('sz'))
                if sz is not None:
                    try: run.font.size = Pt(int(sz.get(w('val'), '24')) / 2)
                    except ValueError: pass
        return p

    def add_table(wtbl):
        rows = wtbl.findall(f'.//{w("tr")}')
        if not rows: return
        cols = max((len(r.findall(f'.//{w("tc")}')) for r in rows), default=1)
        if cols == 0: return
        tbl = doc.add_table(rows=len(rows), cols=cols)
        tbl.style = 'Table Grid'
        for i, wtr in enumerate(rows):
            cells = wtr.findall(f'.//{w("tc")}')
            for j, wtc in enumerate(cells):
                if j < cols:
                    tbl.rows[i].cells[j].text = get_run_text(wtc)

    # Walk body — process w:p and w:tbl in document order
    # Use a recursive walk that finds top-level p/tbl (not nested ones)
    def walk(elem, depth=0):
        for child in elem:
            tag = child.tag
            if tag == w('p'):
                add_para(child)
            elif tag == w('tbl'):
                add_table(child)
            elif tag in (w('body'), w('sect')):
                walk(child, depth+1)
            elif tag not in (w('tr'), w('tc'), w('r'), w('t'),
                             w('pPr'), w('rPr'), w('tPr'), w('tblPr')):
                # recurse into containers (wx:sect etc)
                walk(child, depth+1)

    body = tree.find(w('body'))
    if body is None:
        print("ERROR: no w:body found", file=sys.stderr)
        sys.exit(3)

    walk(body)

    doc.save(output_path)
    size = os.path.getsize(output_path)
    print(f"Converted: {os.path.basename(input_path)} -> {size:,} bytes", file=sys.stderr)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: convert_wordml.py <input.doc> <output.docx>", file=sys.stderr)
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
